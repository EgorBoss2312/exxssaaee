from __future__ import annotations

import io
from pathlib import Path

from docx import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from pypdf import PdfReader


def _docx_body_text(doc: DocxDocument) -> str:
    """Параграфы и ячейки таблиц в порядке следования в теле (не только doc.paragraphs)."""
    parts: list[str] = []
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            p = Paragraph(child, doc)
            if p.text.strip():
                parts.append(p.text)
        elif child.tag == qn("w:tbl"):
            tbl = Table(child, doc)
            for row in tbl.rows:
                for cell in row.cells:
                    for cp in cell.paragraphs:
                        if cp.text.strip():
                            parts.append(cp.text)
    return "\n".join(parts).strip()


def extract_text_from_file(path: Path, mime: str | None) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf" or (mime and "pdf" in mime):
        return _pdf_text(path)
    if suffix in (".docx", ".doc") or (mime and "word" in mime):
        return _docx_text(path)
    if suffix in (".txt", ".md"):
        return path.read_text(encoding="utf-8", errors="replace")
    # fallback: try as text
    try:
        return path.read_text(encoding="utf-8", errors="replace")
    except OSError:
        return ""


def _pdf_text(path: Path) -> str:
    reader = PdfReader(str(path))
    parts: list[str] = []
    for page in reader.pages:
        t = page.extract_text() or ""
        parts.append(t)
    return "\n\n".join(parts).strip()


def _docx_text(path: Path) -> str:
    doc = DocxDocument(str(path))
    return _docx_body_text(doc)


def extract_text_from_bytes(data: bytes, filename: str) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        reader = PdfReader(io.BytesIO(data))
        parts: list[str] = []
        for page in reader.pages:
            parts.append(page.extract_text() or "")
        return "\n\n".join(parts).strip()
    if suffix == ".docx":
        doc = DocxDocument(io.BytesIO(data))
        return _docx_body_text(doc)
    return data.decode("utf-8", errors="replace")
